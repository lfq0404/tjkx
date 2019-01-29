# 从mongo中获取数据，写入word

import pymongo
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

import tjkx.spiders.constants as cons
import tjkx.spiders.control as control


class OutputWord:
    introductions = []
    detail_infos = []

    def get_file(self):
        """
        获取最终文件
        :return:
        """
        self._connection_db()

        # 读取日期范围内的数据，顺序遍历
        for i in self.db.tjkx.find({'public_time': {'$gte': cons.MIN_TIME, '$lte': cons.MAX_TIME}}).sort(
                [('public_time', pymongo.ASCENDING)]):
            # 保存导读
            self.introductions.append(i['introduction'])
            # 保存详情
            self.detail_infos.append(i['details'])

        self._write_word()

    def _connection_db(self):
        """
        链接数据库
        :return:
        """
        connection = pymongo.MongoClient('localhost', 27017)
        self.db = connection['tjkxdb']

    def _write_word(self):
        """
        写入word
        :return:
        """
        # 创建文档对象
        document = Document()

        # 设置默认字体
        document.styles['Normal'].font.name = u'微软雅黑'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 设置文档标题
        document.add_heading(u'上周行业动态回顾', 0)

        # 文档数据范围
        min_date = control.date2chstr(cons.MIN_DATE)
        max_date = control.date2chstr(cons.MAX_DATE)
        document.add_paragraph(u'{0}--{1}'.format(min_date, max_date), style='Subtitle')

        # 一级标题：导读目录
        run = document.add_heading(u'', level=1).add_run(u'上周动态导读目录')
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 导读目录内容（有序列表）
        for introduction in self.introductions:
            # 以分号分隔每一天的导读内容
            intro_temp = introduction.split('；')  # ['四川省人大提案：建立三支酒类发展基金', '“梦之蓝经典咏流传百家姓礼盒”1月22日上线']
            for intro in intro_temp:
                control.del_src_number(intro)
                if intro:
                    run = document.add_paragraph('', style='ListNumber').add_run(intro)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 一级标题：行业动态详情
        run = document.add_heading(u'', level=1).add_run(u'上周行业动态详情')
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 动态内容（有序列表）

        # 保存文档
        document.save('./temp.docx')
